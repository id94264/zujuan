#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
JSON转Word文档工具（本地图片版）
功能：将指定JSON文件转换为格式规范的Word文档，包含表格处理和图片插入
支持SVG图片转换为PNG格式后插入
"""

import json
import os
import sys
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import cairosvg  # 用于SVG转PNG


# 全局常量配置
BASE_IMAGE_HEIGHT_PT = 10.5  # 基准文字高度（五号字）
SMALL_IMAGE_THRESHOLD = 60  # 小图片高度阈值(px)
TEMP_PNG_PATH = "temp_converted.png"  # 临时PNG文件路径


class ImageProcessor:
    """图片处理类"""
    
    def __init__(self, json_file_path=None):
        self.json_file_path = json_file_path
    
    def _convert_svg_to_png(self, svg_path):
        """将SVG图片转换为PNG格式"""
        try:
            # 使用cairosvg将SVG转换为PNG
            cairosvg.svg2png(url=svg_path, write_to=TEMP_PNG_PATH)
            return TEMP_PNG_PATH
        except Exception as e:
            print(f"[ERROR] SVG转PNG失败: {svg_path}, 错误: {str(e)}")
            return None
    
    def _resolve_image_path(self, path):
        """解析图片路径，基于JSON文件所在目录"""
        if not path or not isinstance(path, str):
            return path
        
        path = path.strip()
        
        # 如果已经是绝对路径，直接返回
        if os.path.isabs(path):
            return path
        
        # 如果JSON文件路径存在，基于JSON文件所在目录解析相对路径
        if self.json_file_path and os.path.isfile(self.json_file_path):
            json_dir = os.path.dirname(self.json_file_path)
            resolved_path = os.path.normpath(os.path.join(json_dir, path))
            if os.path.exists(resolved_path):
                return resolved_path
        
        # 否则尝试直接使用路径（可能是相对于当前工作目录）
        return path
    
    def load_local_image(self, path):
        """加载本地图片并返回PIL Image对象"""
        if not path or not isinstance(path, str):
            print(f"[WARNING] 无效的图片路径: {path}")
            return None
        
        path = self._resolve_image_path(path)
        
        try:
            if not os.path.exists(path):
                print(f"[WARNING] 图片文件不存在: {path}")
                return None
                
            # 处理SVG图片
            if path.lower().endswith('.svg'):
                png_path = self._convert_svg_to_png(path)
                if not png_path or not os.path.exists(png_path):
                    return None
                return Image.open(png_path)
            
            return Image.open(path)
        except Exception as e:
            print(f"[ERROR] 加载本地图片失败: {path}, 错误: {str(e)}")
            return None
        finally:
            # 清理临时PNG文件
            if os.path.exists(TEMP_PNG_PATH):
                try:
                    os.remove(TEMP_PNG_PATH)
                except:
                    pass
    
    @staticmethod
    def calculate_image_height(img, is_in_options=False):
        """计算图片应该设置的高度"""
        if not img:
            return Inches(BASE_IMAGE_HEIGHT_PT * 0.8 / 72)
        
        original_height = img.height
        
        if is_in_options:
            # 选项内的图片处理
            if original_height < SMALL_IMAGE_THRESHOLD:
                return Inches(original_height * 0.7 / 72)  # 修改为自身高度的0.5倍
            # return Inches(BASE_IMAGE_HEIGHT_PT * 6 / 72)
            return Inches(original_height * 0.5 / 72)
        
        # 非选项图片处理
        if original_height < SMALL_IMAGE_THRESHOLD:
            return Inches(original_height * 0.7 / 72)  # 修改为自身高度的0.5倍
        # return Inches(BASE_IMAGE_HEIGHT_PT * 6 / 72)
        return Inches(original_height * 0.5 / 72)
    
    @staticmethod
    def should_standalone(img, is_in_options=False):
        """判断图片是否需要单独一行"""
        if not img or is_in_options:
            return False
        
        return img.height >= SMALL_IMAGE_THRESHOLD
    
    def add_to_paragraph(self, paragraph, img_data, is_in_options=False):
        """添加图片到段落"""
        if not isinstance(img_data, dict):
            paragraph.add_run(f"[无效的图片数据]")
            return
        
        img_path = img_data.get('src') or img_data.get('url') or ''
        if not img_path:
            alt_text = img_data.get('alt', '无图片描述')
            paragraph.add_run(f"[图片: {alt_text}]")
            return
        
        img = self.load_local_image(img_path)
        if not img:
            alt_text = img_data.get('alt', '无图片描述')
            paragraph.add_run(f"[图片加载失败: {alt_text}]")
            return
        
        # 计算并设置图片高度
        img_height = self.calculate_image_height(img, is_in_options)
        
        # 将图片保存到内存并插入文档
        try:
            with BytesIO() as img_bytes:
                img_format = 'PNG' if img.mode in ('RGBA', 'LA') else 'JPEG'
                img.save(img_bytes, format=img_format)
                img_bytes.seek(0)
                
                # 在现有段落中添加图片
                run = paragraph.add_run()
                run.add_picture(img_bytes, height=img_height)
                
                # 对于大图片且不在选项中的，添加换行
                if self.should_standalone(img, is_in_options):
                    run.add_break(WD_BREAK.LINE)
                
                return True
        except Exception as e:
            print(f"[ERROR] 插入图片失败: {img_path}, 错误: {str(e)}")
            paragraph.add_run(f"[图片插入失败: {img_data.get('alt', '')}]")
            return False


class DocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self, json_file_path=None):
        self.doc = Document()
        self.image_processor = ImageProcessor(json_file_path)
        self._setup_document()
        self.current_paragraph = None  # 跟踪当前段落
    
    def _setup_document(self):
        """初始化文档格式设置"""
        section = self.doc.sections[0]
        
        # 页面设置
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
        
        # 页眉页脚
        section.header_distance = Cm(1)
        section.footer_distance = Cm(1)
        
        # 无网格设置
        doc_grid = OxmlElement('w:docGrid')
        doc_grid.set(qn('w:type'), 'lines')
        doc_grid.set(qn('w:linePitch'), '0')
        self.doc.element.append(doc_grid)
        
        # 添加样式
        self._add_styles()
    
    def _add_styles(self):
        """创建文档样式"""
        styles = self.doc.styles
        
        # 大标题样式
        title_style = styles.add_style('TitleStyle', 1)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        title_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        
        # 小标题样式
        section_style = styles.add_style('SectionStyle', 1)
        section_style.font.name = 'Times New Roman'
        section_style.font.size = Pt(12)
        section_style.font.bold = True
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        section_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        section_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        
        # 正文样式
        body_style = styles.add_style('BodyStyle', 1)
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(10.5)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        
        # 选项样式
        option_style = styles.add_style('OptionStyle', 1)
        option_style.font.name = 'Times New Roman'
        option_style.font.size = Pt(10.5)
        option_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        option_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        option_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    def _start_new_paragraph(self, style='BodyStyle'):
        """开始一个新的段落"""
        self.current_paragraph = self.doc.add_paragraph(style=style)
        return self.current_paragraph
    
    def _process_content_item(self, item, is_in_options=False):
        """处理单个内容项"""
        if self.current_paragraph is None:
            self._start_new_paragraph('OptionStyle' if is_in_options else 'BodyStyle')
        
        if isinstance(item, str):
            self.current_paragraph.add_run(item)
        elif isinstance(item, dict):
            if item.get('type') == 'image':
                # 处理图片前确保有段落
                if self.current_paragraph is None:
                    self._start_new_paragraph('OptionStyle' if is_in_options else 'BodyStyle')
                
                # 对于需要单独一行的图片，先结束当前段落
                img = self.image_processor.load_local_image(item.get('src', ''))
                if img and self.image_processor.should_standalone(img, is_in_options):
                    if self.current_paragraph.runs:  # 如果当前段落已有内容
                        self._start_new_paragraph('BodyStyle')
                
                self.image_processor.add_to_paragraph(self.current_paragraph, item, is_in_options)
            elif item.get('type') == 'options':
                self._process_options(item.get('content', []))
            elif 'content' in item:
                self._process_content(item['content'], is_in_options)
        elif isinstance(item, list):
            self._process_content(item, is_in_options)
    
    def _process_content(self, content, is_in_options=False):
        """处理内容片段"""
        if isinstance(content, (str, dict)):
            self._process_content_item(content, is_in_options)
        elif isinstance(content, list):
            for item in content:
                self._process_content_item(item, is_in_options)
    
    def _process_options(self, options):
        """处理选项内容"""
        if not isinstance(options, list):
            return
        
        for option in options:
            # 每个选项开始一个新段落
            self._start_new_paragraph('OptionStyle')
            if isinstance(option, str):
                self.current_paragraph.add_run(option)
            else:
                self._process_content(option, is_in_options=True)
    
    def _add_table(self, table_data):
        """添加表格到文档"""
        if not isinstance(table_data, dict) or not isinstance(table_data.get('content'), list):
            return
        
        rows = len(table_data['content'])
        if rows == 0:
            return
        
        cols = max(len(row) for row in table_data['content']) if rows > 0 else 0
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(table_data['content']):
            if not isinstance(row, list):
                continue
                
            for j, cell_content in enumerate(row):
                if j >= cols:
                    continue
                    
                cell = table.cell(i, j)
                # 表格单元格中的内容处理
                self.current_paragraph = cell.paragraphs[0]
                self._process_content(cell_content)
        
        # 表格处理完后重置当前段落
        self.current_paragraph = None
    
    def generate(self, data, output_path):
        """生成Word文档"""
        try:
            # 添加文档标题
            if isinstance(data.get('title'), str):
                self.doc.add_paragraph(data['title'], style='TitleStyle')
            
            # 处理各部分内容
            sections = data.get('sections', [])
            if not isinstance(sections, list):
                sections = []
                
            for section in sections:
                if isinstance(section, dict) and isinstance(section.get('title'), str):
                    self.doc.add_paragraph(section['title'], style='SectionStyle')
                
                questions = section.get('questions', [])
                if not isinstance(questions, list):
                    continue
                    
                for question in questions:
                    # 处理问题内容
                    content = question.get('content', [])
                    if not isinstance(content, list):
                        continue
                    
                    # 重置当前段落
                    self.current_paragraph = None
                    
                    for item in content:
                        if isinstance(item, dict) and item.get('type') == 'table':
                            self._add_table(item)
                            # 表格处理完后重置当前段落
                            self.current_paragraph = None
                        else:
                            self._process_content(item)
            
            # 保存文档
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            self.doc.save(output_path)
            print(f"转换完成！文件已保存至: {os.path.abspath(output_path)}")
            return True
            
        except Exception as e:
            print(f"[CRITICAL] 文档生成失败: {str(e)}", file=sys.stderr)
            return False


def json_to_word(input_path, output_path):
    """主转换函数"""
    try:
        # 验证输入文件
        if not os.path.isfile(input_path):
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
        
        # 读取JSON数据
        with open(input_path, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
            except json.JSONDecodeError as e:
                raise ValueError(f"JSON解析失败: {str(e)}")
        
        # 生成Word文档
        generator = DocumentGenerator(input_path)
        success = generator.generate(data, output_path)
        
        if not success:
            raise RuntimeError("文档生成过程中发生错误")
        
    except Exception as e:
        print(f"[ERROR] 转换过程中发生错误: {str(e)}", file=sys.stderr)
        return False
    
    return True


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("使用方法: python json_to_word.py input.json output.docx", file=sys.stderr)
        sys.exit(1)
    
    success = json_to_word(sys.argv[1], sys.argv[2])
    sys.exit(0 if success else 1)